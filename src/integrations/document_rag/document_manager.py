"""Document manager for handling file uploads and processing."""

import os
import hashlib
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
import pickle

import fitz  # PyMuPDF for PDF processing
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings

logger = logging.getLogger(__name__)

class DocumentManager:
    """Manages document upload, processing, and storage."""
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt', '.md'}
    
    def __init__(self, storage_dir: str = "documents", max_file_size_mb: int = 50):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        self.max_file_size = max_file_size_mb * 1024 * 1024  # Convert to bytes
        
        # Initialize components
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            add_start_index=True
        )
        # Use local sentence-transformers model for embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",  # Fast, lightweight, and works offline
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
        # Load existing documents metadata
        self.metadata_file = self.storage_dir / "documents_metadata.pkl"
        self.documents_metadata = self._load_metadata()
    
    def _load_metadata(self) -> Dict[str, Dict[str, Any]]:
        """Load existing documents metadata."""
        if self.metadata_file.exists():
            try:
                with open(self.metadata_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                logger.warning(f"Failed to load metadata: {e}")
        return {}
    
    def _save_metadata(self):
        """Save documents metadata."""
        try:
            with open(self.metadata_file, 'wb') as f:
                pickle.dump(self.documents_metadata, f)
        except Exception as e:
            logger.error(f"Failed to save metadata: {e}")
    
    def _generate_doc_id(self, filename: str, content: str) -> str:
        """Generate unique document ID based on filename and content."""
        content_hash = hashlib.sha256(content.encode()).hexdigest()[:8]
        return f"{Path(filename).stem}_{content_hash}"
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            raise
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT/MD file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Failed to extract text from TXT {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from supported file formats."""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self._extract_text_from_pdf(file_path)
        elif suffix == '.docx':
            return self._extract_text_from_docx(file_path)
        elif suffix in {'.txt', '.md'}:
            return self._extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def upload_document(self, file_path: str, custom_name: Optional[str] = None) -> Dict[str, Any]:
        """Upload and process a document."""
        source_path = Path(file_path)
        
        # Validate file
        if not source_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if source_path.suffix.lower() not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format. Supported: {self.SUPPORTED_FORMATS}")
        
        if source_path.stat().st_size > self.max_file_size:
            raise ValueError(f"File too large. Max size: {self.max_file_size // (1024*1024)}MB")
        
        # Extract text
        logger.info(f"Processing document: {source_path.name}")
        text_content = self.extract_text(source_path)
        
        if not text_content.strip():
            raise ValueError("Document appears to be empty or text extraction failed")
        
        # Generate document ID
        doc_id = self._generate_doc_id(source_path.name, text_content)
        
        # Check if document already exists
        if doc_id in self.documents_metadata:
            logger.warning(f"Document {doc_id} already exists")
            return self.documents_metadata[doc_id]
        
        # Copy file to storage
        stored_filename = f"{doc_id}{source_path.suffix}"
        stored_path = self.storage_dir / stored_filename
        
        import shutil
        shutil.copy2(source_path, stored_path)
        
        # Create document chunks
        document = Document(
            page_content=text_content,
            metadata={
                "source": stored_filename,
                "doc_id": doc_id,
                "original_filename": source_path.name,
                "file_type": source_path.suffix.lower(),
                "upload_date": datetime.now().isoformat(),
                "file_size": source_path.stat().st_size,
                "custom_name": custom_name or source_path.stem
            }
        )
        
        chunks = self.text_splitter.split_documents([document])
        
        # Create vector store for this document
        vectorstore_path = self.storage_dir / f"{doc_id}_vectorstore"
        vectorstore = FAISS.from_documents(chunks, self.embeddings)
        vectorstore.save_local(str(vectorstore_path))
        
        # Store metadata
        doc_metadata = {
            "doc_id": doc_id,
            "original_filename": source_path.name,
            "custom_name": custom_name or source_path.stem,
            "stored_filename": stored_filename,
            "file_type": source_path.suffix.lower(),
            "upload_date": datetime.now().isoformat(),
            "file_size": source_path.stat().st_size,
            "num_chunks": len(chunks),
            "total_characters": len(text_content),
            "vectorstore_path": str(vectorstore_path)
        }
        
        self.documents_metadata[doc_id] = doc_metadata
        self._save_metadata()
        
        logger.info(f"Document uploaded successfully: {doc_id}")
        return doc_metadata
    
    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get document metadata by ID."""
        return self.documents_metadata.get(doc_id)
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """List all uploaded documents."""
        return list(self.documents_metadata.values())
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its associated files."""
        if doc_id not in self.documents_metadata:
            return False
        
        doc_metadata = self.documents_metadata[doc_id]
        
        try:
            # Delete stored file
            stored_file = self.storage_dir / doc_metadata["stored_filename"]
            if stored_file.exists():
                stored_file.unlink()
            
            # Delete vectorstore directory
            vectorstore_path = Path(doc_metadata["vectorstore_path"])
            if vectorstore_path.exists():
                import shutil
                shutil.rmtree(vectorstore_path)
            
            # Remove from metadata
            del self.documents_metadata[doc_id]
            self._save_metadata()
            
            logger.info(f"Document deleted successfully: {doc_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document {doc_id}: {e}")
            return False
    
    def get_vectorstore(self, doc_id: str) -> Optional[FAISS]:
        """Load vectorstore for a specific document."""
        if doc_id not in self.documents_metadata:
            return None
        
        try:
            vectorstore_path = self.documents_metadata[doc_id]["vectorstore_path"]
            return FAISS.load_local(
                vectorstore_path, 
                self.embeddings,
                allow_dangerous_deserialization=True
            )
        except Exception as e:
            logger.error(f"Failed to load vectorstore for {doc_id}: {e}")
            return None
    
    def search_documents(self, query: str, doc_ids: Optional[List[str]] = None, 
                        k: int = 5) -> List[Dict[str, Any]]:
        """Search across documents for relevant content."""
        results = []
        
        target_docs = doc_ids or list(self.documents_metadata.keys())
        
        for doc_id in target_docs:
            vectorstore = self.get_vectorstore(doc_id)
            if vectorstore:
                try:
                    docs = vectorstore.similarity_search_with_score(query, k=k)
                    for doc, score in docs:
                        results.append({
                            "doc_id": doc_id,
                            "content": doc.page_content,
                            "metadata": doc.metadata,
                            "similarity_score": score,
                            "document_name": self.documents_metadata[doc_id]["custom_name"]
                        })
                except Exception as e:
                    logger.error(f"Search failed for document {doc_id}: {e}")
        
        # Sort by similarity score (lower is better for FAISS)
        results.sort(key=lambda x: x["similarity_score"])
        return results[:k] 